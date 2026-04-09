"""Google Docs/Drive adapter with mock fallback."""

from __future__ import annotations

import json
import re
import textwrap
import unicodedata
import uuid
from abc import ABC, abstractmethod
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

from app.core.config import settings
from app.infrastructure.adapters.google_workspace_auth import build_google_credentials, has_google_oauth_connection
from app.infrastructure.adapters.respiro_template_renderer import RespiroTemplateRenderer

try:
    from pypdf import PdfReader, PdfWriter
except Exception:  # pragma: no cover
    PdfReader = None
    PdfWriter = None


@dataclass(slots=True)
class _SignatureAnchor:
    page_index: int
    x: float
    y: float
    width: float
    height: float
    source: str


class DocsAdapterError(Exception):
    """Raised on document API issues."""


class BaseDocsAdapter(ABC):
    @abstractmethod
    def create_document(self, *, title: str, content: str) -> tuple[str, str]:
        raise NotImplementedError

    @abstractmethod
    def export_docx(
        self,
        *,
        doc_id: str,
        destination_dir: str,
        template_docx_path: str | None = None,
    ) -> tuple[str, str]:
        raise NotImplementedError

    @abstractmethod
    def export_pdf(
        self,
        *,
        doc_id: str,
        destination_dir: str,
        template_pdf_path: str | None = None,
        signature_image_path: str | None = None,
        therapist_name: str | None = None,
    ) -> tuple[str, str]:
        raise NotImplementedError

    def _populate_docx_template(self, *, document: Document, title: str, fields: dict[str, Any]) -> None:
        patient_name = str(fields.get("patient_name", "No registrado"))
        age = str(fields.get("age", "No referido"))
        date = str(fields.get("date", datetime.now().strftime("%d/%m/%Y")))
        sections = fields.get("sections", [])
        section_list = sections if isinstance(sections, list) else []

        if not document.paragraphs:
            document.add_heading(title, 0)

        paragraphs = self._iter_docx_paragraphs(document)
        self._replace_docx_metadata(paragraphs, patient_name=patient_name, age=age, date=date)

        for item in section_list:
            if not isinstance(item, dict):
                continue
            title_text = str(item.get("title", "")).strip()
            body_text = str(item.get("body", "")).strip()
            if not title_text or not body_text:
                continue

            anchor = self._find_section_anchor(paragraphs, title_text)
            body_lines = [line.strip() for line in body_text.splitlines() if line.strip()] or ["No referido"]

            if anchor is None:
                document.add_paragraph("")
                heading = document.add_paragraph(title_text)
                try:
                    heading.style = "Heading 2"
                except Exception:
                    pass
                for line in body_lines:
                    document.add_paragraph(line)
                paragraphs = self._iter_docx_paragraphs(document)
                continue

            cursor = anchor
            for line in body_lines:
                cursor = self._insert_docx_paragraph_after(cursor, line)
            paragraphs = self._iter_docx_paragraphs(document)

    def _replace_docx_metadata(
        self,
        paragraphs: list[Paragraph],
        *,
        patient_name: str,
        age: str,
        date: str,
    ) -> None:
        for paragraph in paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            normalized = self._normalize_docx_text(text)
            if "nombre del paciente" in normalized and ":" in text:
                self._set_paragraph_text(paragraph, f"Nombre del paciente: {patient_name}")
                continue

            if "edad" in normalized and "fecha" in normalized:
                self._set_paragraph_text(paragraph, f"Edad: {age}    Fecha: {date}")
                continue

            if normalized.startswith("edad") and ":" in text:
                self._set_paragraph_text(paragraph, f"Edad: {age}")
                continue

            if "fecha consulta" in normalized and ":" in text:
                self._set_paragraph_text(paragraph, f"Fecha consulta: {date}")
                continue

            if normalized.startswith("fecha") and ":" in text:
                self._set_paragraph_text(paragraph, f"Fecha: {date}")

    def _iter_docx_paragraphs(self, document: Document) -> list[Paragraph]:
        paragraphs = list(document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        return paragraphs

    def _find_section_anchor(self, paragraphs: list[Paragraph], title: str) -> Paragraph | None:
        target = self._normalize_docx_text(title)
        for paragraph in paragraphs:
            normalized = self._normalize_docx_text(paragraph.text)
            if normalized == target or (target and target in normalized):
                return paragraph
        return None

    def _insert_docx_paragraph_after(self, paragraph: Paragraph, text: str) -> Paragraph:
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_paragraph = Paragraph(new_p, paragraph._parent)
        try:
            new_paragraph.style = "Normal"
        except Exception:
            pass
        self._set_paragraph_text(new_paragraph, text)
        return new_paragraph

    def _set_paragraph_text(self, paragraph: Paragraph, text: str) -> None:
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
            return
        paragraph.add_run(text)

    def _normalize_docx_text(self, value: str) -> str:
        lowered = value.lower().replace("-", " ")
        normalized = unicodedata.normalize("NFKD", lowered)
        ascii_only = "".join(ch for ch in normalized if not unicodedata.combining(ch))
        return re.sub(r"\s+", " ", ascii_only).strip()


class MockDocsAdapter(BaseDocsAdapter):
    """Local filesystem-backed docs adapter for demos."""

    def __init__(self) -> None:
        self.base_dir = Path(settings.artifacts_dir)
        self.base_dir.mkdir(parents=True, exist_ok=True)
        templates_dir = Path(__file__).resolve().parents[3] / "assets" / "templates"
        self.default_respiro_docx = templates_dir / "plantilla_historia_clinica_respira_integral_editable_final.docx"
        self.default_respiro_pdf = templates_dir / "plantilla_historia_clinica_respira_integral_of.pdf"
        self.respiro_renderer = RespiroTemplateRenderer()

    def create_document(self, *, title: str, content: str) -> tuple[str, str]:
        doc_id = f"mock-{uuid.uuid4()}"
        payload = {"doc_id": doc_id, "title": title, "content": content}
        path = self.base_dir / f"{doc_id}.gdoc.mock.json"
        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        return doc_id, path.as_uri()

    def export_docx(
        self,
        *,
        doc_id: str,
        destination_dir: str,
        template_docx_path: str | None = None,
    ) -> tuple[str, str]:
        payload = self._read_payload(doc_id)
        destination = Path(destination_dir)
        destination.mkdir(parents=True, exist_ok=True)
        output = destination / f"{doc_id}.docx"

        template = self._resolve_docx_template(template_docx_path)
        fields = self._parse_template_fields(str(payload.get("content", "")))

        if template is not None:
            if self.respiro_renderer.is_respiro_template(template):
                rendered = self.respiro_renderer.render_docx(
                    template_path=str(template),
                    output_path=str(output),
                    fields=fields,
                )
                if rendered:
                    return str(output), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

            document = Document(str(template))
            self._populate_docx_template(document=document, title=str(payload.get("title", "Clinical Draft")), fields=fields)
            document.save(output)
            return str(output), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        document = Document()
        document.add_heading(str(payload.get("title", "Clinical Draft")), 0)
        for paragraph in str(payload.get("content", "")).splitlines():
            document.add_paragraph(paragraph)
        document.save(output)
        return str(output), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def export_pdf(
        self,
        *,
        doc_id: str,
        destination_dir: str,
        template_pdf_path: str | None = None,
        signature_image_path: str | None = None,
        therapist_name: str | None = None,
    ) -> tuple[str, str]:
        payload = self._read_payload(doc_id)
        destination = Path(destination_dir)
        destination.mkdir(parents=True, exist_ok=True)
        output = destination / f"{doc_id}.pdf"
        fields = self._parse_template_fields(str(payload.get("content", "")))

        template = self._resolve_template(template_pdf_path)
        if template is not None and self.respiro_renderer.render_pdf(
            template_path=str(template),
            output_path=str(output),
            fields=fields,
            signature_image_path=signature_image_path,
            therapist_name=therapist_name,
        ):
            return str(output), "application/pdf"

        if template is not None and PdfReader is not None and PdfWriter is not None:
            reader = PdfReader(str(template))
            writer = PdfWriter()
            anchor = self._resolve_signature_anchor(reader)

            for page_index, page in enumerate(reader.pages):
                width = float(page.mediabox.width)
                height = float(page.mediabox.height)
                overlay_buffer = BytesIO()
                overlay = canvas.Canvas(overlay_buffer, pagesize=(width, height))
                overlay.setFillColor(colors.HexColor("#111827"))
                overlay.setFont("Helvetica", 10)

                if page_index == 0:
                    cursor_y = height - 48
                    for line in str(payload.get("content", "")).splitlines():
                        cleaned = line.strip()
                        if not cleaned:
                            cursor_y -= 8
                            continue
                        for wrapped in textwrap.wrap(cleaned, width=95) or [""]:
                            if cursor_y <= 48:
                                break
                            overlay.drawString(40, cursor_y, wrapped)
                            cursor_y -= 12

                if anchor and anchor.page_index == page_index and signature_image_path:
                    self._draw_signature_image(overlay, signature_image_path, anchor)

                overlay.save()
                overlay_buffer.seek(0)
                overlay_page = PdfReader(overlay_buffer).pages[0]
                page.merge_page(overlay_page)
                writer.add_page(page)

            with output.open("wb") as file_handle:
                writer.write(file_handle)
            return str(output), "application/pdf"

        simple = canvas.Canvas(str(output), pagesize=LETTER)
        simple.setFont("Helvetica-Bold", 14)
        simple.drawString(48, 750, str(payload.get("title", "Clinical Draft")))
        simple.setFont("Helvetica", 10)
        cursor_y = 724
        for line in str(payload.get("content", "")).splitlines():
            cleaned = line.strip()
            if not cleaned:
                cursor_y -= 8
                continue
            for wrapped in textwrap.wrap(cleaned, width=92) or [""]:
                if cursor_y <= 48:
                    simple.showPage()
                    simple.setFont("Helvetica", 10)
                    cursor_y = 750
                simple.drawString(48, cursor_y, wrapped)
                cursor_y -= 12
        simple.save()
        return str(output), "application/pdf"

    def _resolve_docx_template(self, profile_template_path: str | None) -> Path | None:
        if profile_template_path and Path(profile_template_path).exists():
            return Path(profile_template_path)
        if self.default_respiro_docx.exists():
            return self.default_respiro_docx
        return None

    def _resolve_template(self, profile_template_path: str | None) -> Path | None:
        if profile_template_path and Path(profile_template_path).exists():
            return Path(profile_template_path)
        if self.default_respiro_pdf.exists():
            return self.default_respiro_pdf
        return None

    def _resolve_signature_anchor(self, reader: PdfReader) -> _SignatureAnchor | None:
        for page_index, page in enumerate(reader.pages):
            annotations = page.get("/Annots") or []
            for annotation_ref in annotations:
                try:
                    annotation = annotation_ref.get_object()
                except Exception:
                    continue
                field_name = str(annotation.get("/T") or "").lower()
                if "responsable" not in field_name and "profesional" not in field_name:
                    continue
                rect = annotation.get("/Rect")
                if not rect or len(rect) != 4:
                    continue
                left, bottom, right, top = [float(value) for value in rect]
                return _SignatureAnchor(
                    page_index=page_index,
                    x=left,
                    y=bottom,
                    width=max(right - left, 24),
                    height=max(top - bottom, 24),
                    source=f"widget:{field_name or 'signature'}",
                )

        for page_index, page in enumerate(reader.pages):
            text = (page.extract_text() or "").lower()
            if "responsable" in text:
                return _SignatureAnchor(
                    page_index=page_index,
                    x=145,
                    y=78,
                    width=180,
                    height=32,
                    source="text:responsable",
                )
            if "profesional tratante" in text:
                return _SignatureAnchor(
                    page_index=page_index,
                    x=145,
                    y=78,
                    width=180,
                    height=32,
                    source="text:profesional_tratante",
                )
        return None

    def _draw_signature_image(self, overlay: canvas.Canvas, signature_image_path: str, anchor: _SignatureAnchor) -> None:
        image_path = Path(signature_image_path)
        if not image_path.exists():
            return
        try:
            image = ImageReader(str(image_path))
            width_raw, height_raw = image.getSize()
            if width_raw <= 0 or height_raw <= 0:
                return
            max_width = max(anchor.width - 4, 20)
            max_height = max(anchor.height - 4, 20)
            scale = min(max_width / width_raw, max_height / height_raw)
            draw_width = width_raw * scale
            draw_height = height_raw * scale
            draw_x = anchor.x + (anchor.width - draw_width) / 2
            draw_y = anchor.y + (anchor.height - draw_height) / 2
            overlay.drawImage(
                image,
                draw_x,
                draw_y,
                width=draw_width,
                height=draw_height,
                preserveAspectRatio=True,
                mask="auto",
            )
        except Exception:
            return

    def _parse_template_fields(self, content: str) -> dict[str, Any]:
        lines = [line.strip() for line in content.splitlines()]

        def is_meaningful(value: str | None) -> bool:
            if value is None:
                return False
            cleaned = value.strip()
            if not cleaned:
                return False
            lowered = cleaned.lower()
            if lowered in {"no referido", "no mencionado", "n/a", "na", "none", "null", "sin dato"}:
                return False
            return not lowered.startswith("no referido")

        def value_for(prefixes: str | list[str], fallback: str) -> str:
            options = [prefixes] if isinstance(prefixes, str) else prefixes
            for prefix in options:
                target = prefix.lower()
                for line in lines:
                    if line.lower().startswith(target):
                        _, _, value = line.partition(":")
                        cleaned = value.strip()
                        if is_meaningful(cleaned):
                            return cleaned
            return fallback

        section_titles = [
            "DATOS PERSONALES DEL PACIENTE",
            "1) DATOS DE IDENTIFICACION",
            "2) ENFERMEDAD ACTUAL (HPI)",
            "3) ANTECEDENTES RELEVANTES",
            "4) SINTOMAS RESPIRATORIOS ACTUALES (CHECKLIST)",
            "5) EVALUACION CLINICA RESPIRATORIA",
            "6) PRUEBAS REALIZADAS EN CONSULTA",
            "7) IMPRESION CLINICA",
            "8) PLAN TERAPEUTICO",
            "RIESGOS Y ALERTAS (PARA REVISION)",
            "RIESGOS Y ALERTAS",
        ]
        footer_prefixes = ("Terapeuta:", "Sesion ID:", "Version del borrador:", "Prompt version:", "Modelo:")
        collected: dict[str, list[str]] = {title: [] for title in section_titles}
        current_section: str | None = None

        for raw_line in lines:
            line = raw_line.strip()
            if not line:
                continue
            matched_title = next((title for title in section_titles if line.upper() == title.upper()), None)
            if matched_title:
                current_section = matched_title
                continue
            if line.startswith(footer_prefixes):
                current_section = None
                continue
            if current_section is not None:
                collected[current_section].append(line)

        address = value_for("- Direccion", "No referido")
        city = value_for("- Ciudad", "No referido")
        if not is_meaningful(city) and is_meaningful(address):
            chunks = [segment.strip() for segment in re.split(r"[,;/]", address) if segment.strip()]
            city = chunks[-1] if chunks else "No referido"

        sections: list[dict[str, str]] = []
        for title in section_titles:
            body_lines = collected.get(title, [])
            if not body_lines:
                continue
            normalized_title = "RIESGOS Y ALERTAS" if title == "RIESGOS Y ALERTAS (PARA REVISION)" else title
            sections.append({"title": normalized_title, "body": "\n".join(body_lines)})

        if not sections:
            sections.append({"title": "PERFIL CLINICO", "body": "Sin informacion suficiente."})

        return {
            "patient_name": value_for(["- Nombre del paciente", "- Nombre"], "No registrado"),
            "age": value_for("- Edad", "No referido"),
            "date": value_for("- Fecha consulta", datetime.now().strftime("%d/%m/%Y")),
            "birth_date": value_for("- Fecha de nacimiento", "No referido"),
            "phone": value_for("- Telefono", "No referido"),
            "city": city,
            "email": value_for("- Email", "No referido"),
            "identification": value_for("- Identificacion", "No referido"),
            "therapist_name": value_for("Terapeuta", "Profesional tratante"),
            "sections": sections,
        }

    def _read_payload(self, doc_id: str) -> dict[str, str]:
        source = self.base_dir / f"{doc_id}.gdoc.mock.json"
        if not source.exists():
            raise DocsAdapterError(f"Mock document not found: {doc_id}")
        return json.loads(source.read_text(encoding="utf-8"))


class GoogleDocsAdapter(BaseDocsAdapter):
    """Credentialed Google Docs adapter using domain-wide delegation."""

    def __init__(self, *, therapist = None, impersonated_user: str | None = None) -> None:
        from googleapiclient.discovery import build

        scopes = [
            "https://www.googleapis.com/auth/documents",
            "https://www.googleapis.com/auth/drive",
        ]
        try:
            creds = build_google_credentials(scopes=scopes, therapist=therapist, impersonated_user=impersonated_user)
        except ValueError as exc:
            raise DocsAdapterError(str(exc)) from exc

        self.docs_service = build("docs", "v1", credentials=creds, cache_discovery=False)
        self.drive_service = build("drive", "v3", credentials=creds, cache_discovery=False)

    def create_document(self, *, title: str, content: str) -> tuple[str, str]:
        doc = self.docs_service.documents().create(body={"title": title}).execute()
        doc_id = doc["documentId"]
        self.docs_service.documents().batchUpdate(
            documentId=doc_id,
            body={
                "requests": [
                    {
                        "insertText": {
                            "location": {"index": 1},
                            "text": content,
                        }
                    }
                ]
            },
        ).execute()

        if settings.google_docs_output_folder_id:
            self.drive_service.files().update(
                fileId=doc_id,
                addParents=settings.google_docs_output_folder_id,
                fields="id, parents",
            ).execute()

        return doc_id, f"https://docs.google.com/document/d/{doc_id}/edit"

    def export_docx(
        self,
        *,
        doc_id: str,
        destination_dir: str,
        template_docx_path: str | None = None,
    ) -> tuple[str, str]:
        _ = template_docx_path
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return self._export_file(doc_id=doc_id, destination_dir=destination_dir, suffix="docx", mime_type=mime_type)

    def export_pdf(
        self,
        *,
        doc_id: str,
        destination_dir: str,
        template_pdf_path: str | None = None,
        signature_image_path: str | None = None,
        therapist_name: str | None = None,
    ) -> tuple[str, str]:
        _ = template_pdf_path
        _ = signature_image_path
        _ = therapist_name
        return self._export_file(doc_id=doc_id, destination_dir=destination_dir, suffix="pdf", mime_type="application/pdf")

    def _export_file(self, *, doc_id: str, destination_dir: str, suffix: str, mime_type: str) -> tuple[str, str]:
        request = self.drive_service.files().export_media(fileId=doc_id, mimeType=mime_type)
        destination = Path(destination_dir)
        destination.mkdir(parents=True, exist_ok=True)
        output = destination / f"{doc_id}.{suffix}"
        output.write_bytes(request.execute())
        return str(output), mime_type


def build_docs_adapter(*, therapist = None, impersonated_user: str | None = None) -> BaseDocsAdapter:
    if settings.use_mock_google:
        return MockDocsAdapter()
    if not settings.google_service_account_file and not has_google_oauth_connection(therapist):
        return MockDocsAdapter()
    return GoogleDocsAdapter(therapist=therapist, impersonated_user=impersonated_user)
