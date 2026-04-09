from pathlib import Path

from docx import Document
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.application.services.pipeline_service import PipelineService
from app.application.services.prompt_registry import PromptRegistry
from app.domain.models import Session as SessionModel
from tests.helpers import seed_therapist_and_session


def test_draft_profile_adapts_to_uploaded_template_order(db_session: Session, tmp_path: Path) -> None:
    _, session_id = seed_therapist_and_session(db_session)
    session = db_session.scalar(select(SessionModel).where(SessionModel.id == session_id))
    assert session is not None

    template_path = tmp_path / "custom-template.docx"
    document = Document()
    document.add_paragraph("Plan terapeutico")
    document.add_paragraph("Datos de identificacion")
    document.add_paragraph("Sintomas respiratorios actuales")
    document.add_paragraph("Impresion clinica")
    document.save(template_path)

    session.therapist.template_docx_path = str(template_path)
    db_session.add(session.therapist)

    registry = PromptRegistry(db_session)
    registry.ensure_seeded(["clinical_draft"])

    pipeline = PipelineService(db_session)
    pipeline.ingest_from_meet(session)
    draft = pipeline.generate_draft(session)

    assert draft.clinical_profile_text.startswith("8) Plan terapeutico")
    assert "1) Datos de identificacion" in draft.clinical_profile_text
    assert "4) Sintomas respiratorios actuales (Checklist)" in draft.clinical_profile_text
    assert "7) Impresion clinica" in draft.clinical_profile_text
